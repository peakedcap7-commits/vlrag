from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DESKTOP = Path.home() / "Desktop"
OUT = DESKTOP / "AI私人导购应用项目计划书-技术评审版.docx"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_style(table, header_fill="E8EEF5"):
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.margin_top = 80
            cell.margin_bottom = 80
            cell.margin_left = 120
            cell.margin_right = 120
            if row_idx == 0:
                set_cell_shading(cell, header_fill)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if level == 1:
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        else:
            run.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
    return p


def add_para(doc: Document, text: str = "", bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        rest = text[len(bold_prefix):]
        r2 = p.add_run(rest)
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for run in runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_style(table)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        if widths:
            table.rows[0].cells[i].width = Inches(widths[i])
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value)
            if widths:
                row.cells[i].width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)

    # 封面
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("AI 私人导购应用项目计划书")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("技术评审版 · 多模态 RAG + LangGraph 导购编排 + OpenClaw 风格低权限前端")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_table(
        doc,
        ["文档属性", "内容"],
        [
            ["项目名称", "AI 私人导购应用（ShoppingQnA 升级版）"],
            ["目标读者", "技术负责人、后端工程师、前端工程师、算法工程师、产品负责人"],
            ["文档目的", "明确业务需求、分层架构、技术栈、数据结构、向量存储和数据流，为后续开发评审提供依据"],
            ["建议阶段", "v0.3-v0.6：从本地多模态 RAG Demo 升级为可演示 Web 产品"],
        ],
        [1.5, 4.9],
    )

    add_heading(doc, "1. 项目定位与建设目标")
    add_para(doc, "本项目定位为面向服饰电商场景的 AI 私人导购应用。系统不只是回答商品问题，而是通过多模态理解、用户意图解析、商品检索、偏好记忆、成套穿搭推荐和推荐解释，帮助用户完成购物决策。")
    add_para(doc, "核心目标：将当前多模态 RAG 商品问答助手升级为具备场景理解、用户偏好沉淀、商品组合推荐和可解释推荐能力的 AI 导购系统。", "核心目标：")

    add_heading(doc, "2. 业务需求说明")
    add_table(
        doc,
        ["业务模块", "业务说明", "AI 能力", "优先级"],
        [
            ["自然语言导购", "用户用中文描述场景、预算、颜色、风格、禁忌，系统生成推荐方案", "意图识别、结构化抽取、RAG 生成", "P0"],
            ["图片找相似款", "用户上传服饰图片，系统找相似风格、相似颜色、相似品类商品", "视觉编码、图片向量检索、多模态理解", "P0"],
            ["图文混合检索", "用户上传图片并补充要求，如“类似这件但低调一点”", "图片检索 + 文本检索 + Hybrid 融合", "P0"],
            ["成套穿搭推荐", "按场景推荐外套、上装、下装、鞋子等组合方案", "类别召回、组合评分、推荐解释", "P1"],
            ["用户偏好记忆", "记住用户长期偏好，如黑白灰、低调、不喜欢大 Logo", "对话总结、偏好抽取、记忆更新", "P1"],
            ["商品对比", "对比两件或多件商品，解释适用场景和取舍", "多商品上下文对比、维度评分", "P2"],
            ["主动反问", "需求不完整时主动询问预算、场景、颜色等关键信息", "需求完整性判断、对话状态机", "P2"],
        ],
        [1.05, 2.15, 2.2, 0.55],
    )

    add_heading(doc, "3. 分层系统架构")
    add_para(doc, "系统采用分层架构：前端体验层负责交互，API 层负责安全入口，调度层负责编排 AI 导购流程，模型层封装大模型与嵌入模型，检索层负责多路召回与排序，存储层负责商品、会话、偏好和向量数据。")
    add_table(
        doc,
        ["层级", "职责", "核心组件", "边界说明"],
        [
            ["前端体验层", "聊天、图片上传、商品卡片、搭配方案、过程展示", "React / Next.js、Tailwind、SSE 客户端", "只做展示和请求，不保存模型 Key，不执行工具"],
            ["API 接入层", "鉴权、限流、路由、请求校验、流式响应", "FastAPI、Pydantic、SSE", "所有前端请求必须经过 API 层"],
            ["调度层", "意图解析、分支决策、反问、召回、组合、质检、记忆更新", "LangGraph、业务服务 Service", "只编排流程，不直接写复杂业务细节"],
            ["模型层", "LLM、多模态模型、Embedding、Reranker 封装", "qwen-max、qwen-vl-max、text-embedding-v3、OpenCLIP/Chinese-CLIP", "统一封装模型调用、重试、超时和日志"],
            ["检索层", "文本检索、图片检索、Hybrid 融合、标签过滤、重排序", "Chroma/Qdrant、RRF、LLM Rerank", "检索服务向调度层返回候选商品，不生成最终回答"],
            ["存储层", "商品、标签、用户、会话、偏好、推荐日志、向量索引", "SQLite/PostgreSQL、Redis、对象存储、向量库", "业务数据和向量数据分开管理"],
        ],
        [0.95, 1.75, 1.85, 1.85],
    )

    add_heading(doc, "4. 架构说明与数据流交互逻辑")
    add_para(doc, "在线导购流程以用户会话为核心。每次请求进入后端后，系统先解析用户意图，再合并用户画像，随后进入检索与推荐编排。对于信息不完整的请求，系统应优先反问，而不是强行推荐。")
    add_table(
        doc,
        ["步骤", "处理节点", "输入", "输出"],
        [
            ["1", "输入预处理", "文本、图片、session_id", "标准化文本、图片路径、会话上下文"],
            ["2", "意图解析", "用户输入 + 历史上下文", "intent、scene、style、colors、budget、categories、avoid"],
            ["3", "偏好合并", "结构化需求 + 用户画像", "本轮推荐约束条件"],
            ["4", "多路召回", "约束条件、文本 query、图片向量", "候选商品池"],
            ["5", "过滤与重排序", "候选商品池、标签、偏好", "Top-K 商品或候选组合"],
            ["6", "穿搭组合", "不同品类候选商品", "1-3 套搭配方案"],
            ["7", "约束检查", "推荐方案 + 用户需求", "通过 / 重新检索 / 重新组合 / 反问"],
            ["8", "回答生成", "最终商品方案和解释要点", "流式回答 + 商品卡片数据"],
            ["9", "记忆更新", "用户反馈、显式偏好、隐式行为", "更新 user_preference、recommendation_log"],
        ],
        [0.45, 1.2, 2.15, 2.6],
    )

    add_heading(doc, "5. LangGraph 编排设计")
    add_para(doc, "LangGraph 不建议覆盖整个项目，只用于 AI 导购决策编排。数据加载、向量入库、数据库 CRUD、前端 UI 和普通 API 路由仍使用常规工程范式实现。")
    add_table(
        doc,
        ["节点", "职责", "是否可循环", "失败处理"],
        [
            ["parse_intent", "解析意图、槽位、场景和需求完整度", "否", "解析失败时降级为普通商品搜索"],
            ["load_profile", "加载用户长期偏好和当前会话短期记忆", "否", "读取失败时使用空画像"],
            ["clarify_or_continue", "判断是否反问、检索或更新偏好", "是", "最多反问 2 次"],
            ["retrieve_candidates", "执行文本、图片或 Hybrid 检索", "是", "结果不足时改写 query 后重检索"],
            ["compose_outfit", "按品类组合穿搭方案", "是", "缺品类时单独补召回"],
            ["check_constraints", "检查颜色、场景、预算、禁忌是否满足", "是", "不满足则重新组合，最多 2 轮"],
            ["generate_answer", "生成解释、商品卡片、下一步建议", "否", "失败时返回模板化回答"],
            ["update_memory", "抽取并更新用户偏好", "否", "写入失败不影响本次推荐"],
        ],
        [1.15, 2.1, 1.0, 2.15],
    )

    add_heading(doc, "6. 完整技术栈")
    add_table(
        doc,
        ["类别", "推荐技术", "用途", "阶段建议"],
        [
            ["大模型框架", "LangChain + LangGraph", "模型调用封装、导购流程图编排", "v0.3 起"],
            ["大模型", "qwen-max / qwen-turbo", "意图解析、推荐解释、偏好总结、轻量任务", "当前可用"],
            ["多模态模型", "qwen-vl-max", "图片理解、结构化视觉标签生成", "当前可用"],
            ["文本向量", "text-embedding-v3", "商品文本、用户偏好、查询语义向量", "当前可用"],
            ["图片向量", "OpenCLIP，后续可评估 Chinese-CLIP", "图片相似检索、图文检索", "当前可用"],
            ["向量库", "Chroma 起步，生产可换 Qdrant / Milvus", "文本向量和图片向量存储", "当前 Chroma"],
            ["后端", "FastAPI + Pydantic", "API、SSE、文件上传、业务服务", "下一步建设"],
            ["前端", "Next.js / React + Tailwind + shadcn/ui", "OpenClaw 风格低权限导购前端", "下一步建设"],
            ["缓存", "Redis", "会话缓存、短期记忆、任务状态", "v0.4"],
            ["数据库", "SQLite 起步，生产 PostgreSQL", "商品、用户、偏好、推荐日志", "v0.3"],
            ["容器", "Docker Compose", "后端、前端、Redis、数据库、向量库本地编排", "v0.4"],
            ["对象存储", "本地目录起步，生产 MinIO / OSS", "商品图片、上传图片、生成文件", "v0.4"],
            ["测试", "pytest + httpx + Playwright", "单元、接口、端到端验证", "v0.3"],
            ["观测", "结构化日志 + 请求 ID + 调用耗时统计", "模型、检索、推荐链路排查", "v0.3"],
        ],
        [1.05, 1.7, 2.45, 1.05],
    )

    add_heading(doc, "7. 业务数据表设计")
    add_para(doc, "数据库建议先使用 SQLite 实现，表结构保持可迁移到 PostgreSQL。向量数据不直接放业务表，业务表保存向量库 collection 名称、document_id 和源字段。")

    tables = {
        "users 用户表": (
            ["字段", "类型", "索引", "说明"],
            [
                ["id", "TEXT", "主键", "用户唯一标识，匿名用户可使用 uuid"],
                ["nickname", "TEXT", "普通索引", "用户昵称，可为空"],
                ["created_at", "DATETIME", "普通索引", "创建时间"],
                ["updated_at", "DATETIME", "无", "更新时间"],
            ],
        ),
        "sessions 会话表": (
            ["字段", "类型", "索引", "说明"],
            [
                ["id", "TEXT", "主键", "会话 ID"],
                ["user_id", "TEXT", "外键索引 users.id", "所属用户"],
                ["title", "TEXT", "无", "会话标题"],
                ["status", "TEXT", "普通索引", "active / archived"],
                ["created_at", "DATETIME", "普通索引", "创建时间"],
                ["updated_at", "DATETIME", "普通索引", "更新时间"],
            ],
        ),
        "messages 消息表": (
            ["字段", "类型", "索引", "说明"],
            [
                ["id", "TEXT", "主键", "消息 ID"],
                ["session_id", "TEXT", "外键索引 sessions.id", "所属会话"],
                ["role", "TEXT", "普通索引", "user / assistant / system / tool"],
                ["content", "TEXT", "无", "消息正文"],
                ["image_path", "TEXT", "无", "用户上传图片路径，可为空"],
                ["metadata_json", "TEXT", "无", "检索过程、模型参数等扩展信息"],
                ["created_at", "DATETIME", "普通索引", "创建时间"],
            ],
        ),
        "products 商品表": (
            ["字段", "类型", "索引", "说明"],
            [
                ["id", "INTEGER", "主键", "商品 ID，对应 product_id"],
                ["source_text", "TEXT", "无", "原始数据集 text 字段"],
                ["category", "TEXT", "普通索引", "商品类别，如 outer / bottom / shoes"],
                ["name", "TEXT", "普通索引", "商品名称"],
                ["summary", "TEXT", "无", "原始英文摘要"],
                ["image_path", "TEXT", "唯一索引", "本地图片路径或对象存储地址"],
                ["image_summary", "TEXT", "无", "多模态模型生成的视觉描述"],
                ["description", "TEXT", "无", "营销文案或推荐用描述"],
                ["status", "TEXT", "普通索引", "active / inactive"],
                ["created_at", "DATETIME", "普通索引", "创建时间"],
                ["updated_at", "DATETIME", "无", "更新时间"],
            ],
        ),
        "product_tags 商品标签表": (
            ["字段", "类型", "索引", "说明"],
            [
                ["id", "INTEGER", "主键", "标签记录 ID"],
                ["product_id", "INTEGER", "外键索引 products.id", "关联商品"],
                ["tag_type", "TEXT", "联合索引 product_id, tag_type", "color / style / season / scene / material / feature"],
                ["tag_value", "TEXT", "联合索引 tag_type, tag_value", "标准化标签值"],
                ["source", "TEXT", "普通索引", "vl_model / llm_normalizer / manual"],
                ["confidence", "REAL", "无", "标签置信度，0-1"],
            ],
        ),
        "user_preferences 用户偏好表": (
            ["字段", "类型", "索引", "说明"],
            [
                ["id", "TEXT", "主键", "偏好记录 ID"],
                ["user_id", "TEXT", "外键索引 users.id", "所属用户"],
                ["preference_type", "TEXT", "联合索引 user_id, preference_type", "preferred_color / avoid_style / budget 等"],
                ["preference_value", "TEXT", "联合索引 preference_type, preference_value", "偏好值"],
                ["weight", "REAL", "普通索引", "偏好强度"],
                ["source_message_id", "TEXT", "外键 messages.id", "来源消息"],
                ["created_at", "DATETIME", "普通索引", "创建时间"],
                ["updated_at", "DATETIME", "无", "更新时间"],
            ],
        ),
        "recommendation_logs 推荐日志表": (
            ["字段", "类型", "索引", "说明"],
            [
                ["id", "TEXT", "主键", "推荐日志 ID"],
                ["session_id", "TEXT", "外键索引 sessions.id", "所属会话"],
                ["user_id", "TEXT", "外键索引 users.id", "所属用户"],
                ["intent_json", "TEXT", "无", "结构化意图"],
                ["candidate_ids_json", "TEXT", "无", "候选商品 ID 列表"],
                ["result_json", "TEXT", "无", "最终推荐商品和搭配方案"],
                ["strategy", "TEXT", "普通索引", "text / image / hybrid / outfit"],
                ["latency_ms", "INTEGER", "普通索引", "总耗时"],
                ["created_at", "DATETIME", "普通索引", "创建时间"],
            ],
        ),
        "feedback 用户反馈表": (
            ["字段", "类型", "索引", "说明"],
            [
                ["id", "TEXT", "主键", "反馈 ID"],
                ["recommendation_id", "TEXT", "外键 recommendation_logs.id", "关联推荐"],
                ["user_id", "TEXT", "外键索引 users.id", "反馈用户"],
                ["product_id", "INTEGER", "外键索引 products.id", "关联商品，可为空"],
                ["feedback_type", "TEXT", "普通索引", "like / dislike / click / save / reject"],
                ["comment", "TEXT", "无", "用户文字反馈"],
                ["created_at", "DATETIME", "普通索引", "创建时间"],
            ],
        ),
    }

    for title, (headers, rows) in tables.items():
        add_heading(doc, title, level=2)
        add_table(doc, headers, rows, [1.25, 1.0, 1.65, 2.5])

    add_heading(doc, "8. 向量存储结构设计")
    add_table(
        doc,
        ["Collection", "向量模型", "维度", "Document 内容", "Metadata", "用途"],
        [
            ["product_text_vectors", "text-embedding-v3", "1024", "description / image_summary / summary / normalized_tags", "product_id, field, category, colors, styles, scenes, seasons", "文本语义检索、偏好匹配、RAG 上下文"],
            ["product_image_vectors", "OpenCLIP / Chinese-CLIP", "512", "image_path", "product_id, category, name, colors, styles", "上传图片找相似款、图文混合召回"],
            ["user_preference_vectors", "text-embedding-v3", "1024", "用户偏好摘要", "user_id, preference_type, updated_at", "长期偏好召回、个性化推荐"],
            ["session_memory_vectors", "text-embedding-v3", "1024", "会话摘要、关键约束", "session_id, user_id, turn_range", "长会话记忆检索"],
        ],
        [1.35, 1.25, 0.5, 1.55, 2.15, 1.55],
    )

    add_heading(doc, "9. 模型调用与安全边界")
    add_bullets(
        doc,
        [
            "前端不得保存 DASHSCOPE_API_KEY，不得直接调用大模型 API。",
            "LLM 只能输出结构化意图或推荐内容，不能直接执行系统命令、修改文件或访问任意工具。",
            "后端工具调用必须采用白名单机制，例如 search_products、recommend_outfit、analyze_image、update_preference。",
            "所有模型调用记录 request_id、model、latency_ms、token_usage、error_code，便于排查和成本统计。",
            "图片上传需限制大小、格式和存储目录，避免任意路径写入。",
        ]
    )

    add_heading(doc, "10. 里程碑规划")
    add_table(
        doc,
        ["版本", "目标", "核心交付", "验收标准"],
        [
            ["v0.3", "AI 导购基础版", "FastAPI、意图解析、商品标签标准化、文本/图片/Hybrid 推荐", "文本和图片查询可稳定返回商品卡片"],
            ["v0.4", "偏好记忆版", "用户偏好表、会话表、偏好更新节点、Redis 缓存", "多轮对话能记住颜色、风格、禁忌偏好"],
            ["v0.5", "成套搭配版", "OutfitComposer、组合评分、约束检查、反问机制", "可返回至少 1 套完整穿搭方案并解释原因"],
            ["v0.6", "Web 产品版", "OpenClaw 风格低权限前端、SSE 流式输出、商品详情侧栏", "前后端完整可演示，用户无法通过对话修改系统功能"],
        ],
        [0.75, 1.45, 2.45, 1.85],
    )

    add_heading(doc, "11. 技术风险与应对")
    add_table(
        doc,
        ["风险", "影响", "应对策略"],
        [
            ["模型输出不稳定", "结构化解析失败、推荐理由漂移", "使用 JSON Schema 校验、失败重试、模板化降级"],
            ["OpenCLIP 中文检索弱", "中文文本查图片效果差", "中文先翻译英文，后续评估 Chinese-CLIP"],
            ["向量检索相关性不足", "推荐不准", "引入标签过滤、RRF、Reranker、人工评估集"],
            ["偏好记忆污染", "错误记住用户偏好", "仅记录明确偏好，允许用户查看和删除偏好"],
            ["前端越权风险", "用户通过对话触发危险能力", "低权限业务 API、工具白名单、禁止命令执行和文件编辑接口"],
            ["工程结构臃肿", "后续迭代困难", "采用分层架构，LangGraph 只负责导购编排，数据处理与服务层独立"],
        ],
        [1.65, 1.6, 3.15],
    )

    add_heading(doc, "12. 结论")
    add_para(doc, "本项目建议从现有多模态 RAG 商品问答原型升级为 AI 私人导购应用。技术路线不应追求全项目 LangGraph 化，而应采用分层架构：FastAPI 提供安全接口，LangGraph 编排导购决策，服务层封装检索、商品、偏好和搭配能力，存储层分别管理业务数据和向量数据。")
    add_para(doc, "评审重点：先完成业务数据结构、偏好记忆、意图解析和低权限前端，再逐步增加成套穿搭、商品对比和平替推荐。", "评审重点：")

    # 页脚
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("AI 私人导购应用项目计划书 · 技术评审版")
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_doc()
    print(path)
