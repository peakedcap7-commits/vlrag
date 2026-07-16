from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path.home() / "Desktop" / "AI私人导购应用项目计划书-技术评审版.docx"
FALLBACK_OUT_PATH = Path.home() / "Desktop" / "AI私人导购应用项目计划书-技术评审版-v2.docx"


def set_run_font(run, size: float = 10.5, bold: bool = False, color: str | None = None):
    """统一设置中英文字体，避免 Word 默认字体漂移。"""
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str):
    """设置表头背景色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, bold: bool = False):
    """写入单元格内容并设置紧凑但可读的样式。"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run_font(run, size=8.6, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph(doc: Document, text: str, prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if prefix and text.startswith(prefix):
        r1 = p.add_run(prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(prefix):])
        set_run_font(r2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        set_run_font(
            run,
            size={1: 16, 2: 13, 3: 12}.get(level, 11),
            bold=True,
            color="2E74B5" if level == 1 else "1F4D78",
        )
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(item)
        set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "E8EEF5")
        set_cell(cell, header, bold=True)
        if widths:
            cell.width = Inches(widths[idx])

    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            set_cell(row.cells[idx], value)
            if widths:
                row.cells[idx].width = Inches(widths[idx])

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Microsoft YaHei"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return doc


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("AI 私人导购应用项目计划书")
    set_run_font(run, size=24, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("技术评审版 · 多模态 RAG + 图数据库 + 用户记忆 + LangGraph 受控编排")
    set_run_font(run, size=12, color="555555")

    add_table(
        doc,
        ["文档属性", "内容"],
        [
            ["项目名称", "AI 私人导购应用（ShoppingQnA 升级版）"],
            ["核心定位", "面向服饰电商场景的多模态 AI 导购、搭配推荐与个性化购物决策助手"],
            ["一期强依赖", "FastAPI、LangGraph、Chroma、Neo4j、SQLite/PostgreSQL、Redis、百炼模型、OpenCLIP"],
            ["文档目的", "明确业务需求、分层架构、数据流、数据表、图谱结构、向量存储和评审边界"],
        ],
        [1.5, 4.8],
    )


def add_business_requirements(doc: Document):
    add_heading(doc, "1. 业务需求说明")
    add_paragraph(
        doc,
        "本项目从“商品问答搜索”升级为“AI 私人导购”。系统需要理解用户场景、风格、预算、颜色、禁忌和历史偏好，并结合商品文本、图片、标签、关系图谱和用户记忆，生成可解释的单品推荐、成套穿搭、相似款和平替方案。",
    )
    add_table(
        doc,
        ["业务能力", "用户场景", "系统输出", "优先级"],
        [
            ["自然语言导购", "用户描述场景、预算、颜色、风格，如“冬天通勤，黑灰低调”", "结构化需求、商品推荐、推荐理由", "P0"],
            ["图片找相似款", "用户上传一张衣服图片", "相似款、同风格款、同类别款", "P0"],
            ["图文混合推荐", "用户上传图片并补充“类似这件但更低调”", "图片相似 + 文本约束的混合结果", "P0"],
            ["成套穿搭推荐", "用户要求“一套适合出差/约会/户外的搭配”", "外套/上装/下装/鞋子组合方案", "P1"],
            ["用户偏好记忆", "用户表达“不喜欢花哨”“以后优先黑白灰”", "短期会话记忆 + 长期偏好画像", "P1"],
            ["图谱关系推荐", "用户问“这件夹克配什么裤子”“有没有平替”", "搭配关系、相似关系、替代关系解释", "P1"],
            ["商品对比", "用户比较两件商品哪个更适合通勤", "维度化对比、优缺点、适用场景", "P2"],
            ["主动反问", "用户需求缺少关键条件", "预算/场景/颜色/品类等澄清问题", "P2"],
        ],
        [1.15, 2.0, 2.3, 0.55],
    )


def add_layered_architecture(doc: Document):
    add_heading(doc, "2. 分层系统架构")
    add_paragraph(
        doc,
        "系统采用分层架构。图数据库从一期纳入核心链路，用于商品关系、搭配关系、平替关系、场景关系和用户偏好关系建模；向量库用于语义召回；关系型数据库用于业务事实存储；Redis 用于会话和短期状态缓存。",
    )
    add_table(
        doc,
        ["层级", "职责", "核心组件", "一期是否纳入"],
        [
            ["前端体验层", "聊天、图片上传、商品卡片、搭配方案、关系解释、偏好面板", "Next.js / React / Tailwind / shadcn/ui", "是"],
            ["API 接入层", "鉴权、限流、参数校验、SSE 流式响应、文件上传", "FastAPI / Pydantic / SSE", "是"],
            ["调度层", "意图解析、反问、检索编排、图谱扩展、搭配组合、约束检查、记忆更新", "LangGraph + Service 层", "是"],
            ["模型层", "LLM、多模态理解、Embedding、Rerank、标签归一化", "qwen-max / qwen-vl-max / qwen-turbo / text-embedding-v3 / OpenCLIP", "是"],
            ["检索层", "文本检索、图片检索、标签过滤、图谱邻居扩展、RRF 融合、重排序", "Chroma + Neo4j + RRF + LLM Rerank", "是"],
            ["存储层", "业务表、会话、偏好、商品关系、向量索引、图片文件", "SQLite/PostgreSQL + Neo4j + Chroma + Redis + 本地/OSS", "是"],
            ["观测层", "请求日志、模型耗时、检索命中、推荐链路、错误追踪", "结构化日志 + request_id + 指标统计", "是"],
        ],
        [1.0, 2.05, 2.25, 0.7],
    )


def add_tech_stack(doc: Document):
    add_heading(doc, "3. 完整技术栈")
    add_table(
        doc,
        ["类别", "技术选型", "用途", "说明"],
        [
            ["大模型框架", "LangChain + LangGraph", "模型调用封装、受控 Agentic Workflow 编排", "LangGraph 只管在线导购决策，不管离线入库"],
            ["大语言模型", "qwen-max / qwen-turbo", "意图解析、偏好总结、推荐解释、轻量归一化", "百炼 API"],
            ["多模态模型", "qwen-vl-max", "图片理解、视觉标签、上传图片描述", "百炼 API"],
            ["文本向量", "text-embedding-v3", "商品文本、用户记忆、查询语义向量", "1024 维"],
            ["图片向量", "OpenCLIP，后续评估 Chinese-CLIP", "商品图片和上传图片向量化", "512 维"],
            ["向量库", "Chroma 一期，Qdrant/Milvus 可替换", "文本向量、图片向量、用户记忆向量", "本地 Demo 优先 Chroma"],
            ["图数据库", "Neo4j Community", "商品关系、搭配关系、平替关系、用户偏好关系", "一期纳入核心链路"],
            ["关系数据库", "SQLite 一期，PostgreSQL 生产", "商品、会话、消息、推荐日志、偏好事实", "保持可迁移"],
            ["缓存", "Redis", "短期会话、LangGraph checkpoint、任务状态、热点商品", "一期可先本地 Redis"],
            ["后端", "FastAPI + Pydantic", "API、SSE、图片上传、业务服务", "前端不得直连模型"],
            ["前端", "Next.js / React / Tailwind / shadcn/ui", "OpenClaw 风格低权限导购界面", "只展示和发业务请求"],
            ["容器", "Docker Compose", "后端、前端、Redis、Neo4j、数据库、向量库统一启动", "评审和演示必备"],
            ["测试", "pytest + httpx + Playwright", "单元、接口、端到端、前端交互测试", "v0.3 起补齐"],
        ],
        [1.05, 1.7, 2.1, 1.55],
    )


def add_data_flow(doc: Document):
    add_heading(doc, "4. 架构说明与数据流交互逻辑")
    add_heading(doc, "4.1 商品离线入库流程", level=2)
    add_table(
        doc,
        ["步骤", "处理动作", "输入", "输出/落库"],
        [
            ["1", "加载商品数据", "Kream 商品图片、text 字段", "products 原始记录"],
            ["2", "图片落盘", "PIL Image", "image_path"],
            ["3", "多模态增强", "图片 + 商品名 + 原始摘要", "image_summary、视觉 tags"],
            ["4", "标签标准化", "自由标签 + 描述文本", "标准颜色、风格、季节、场景、材质、特征"],
            ["5", "文本向量化", "description / image_summary / summary / tags", "product_text_vectors"],
            ["6", "图片向量化", "image_path", "product_image_vectors"],
            ["7", "图谱构建", "商品标签、类别、风格、LLM 搭配判断", "Neo4j Product/Style/Scene/Category 节点与关系"],
            ["8", "业务表写入", "增强后的商品记录", "products、product_tags、product_relations"],
        ],
        [0.45, 1.35, 2.0, 2.8],
    )

    add_heading(doc, "4.2 在线导购推荐流程", level=2)
    add_table(
        doc,
        ["步骤", "节点", "说明", "结果"],
        [
            ["1", "输入预处理", "接收文本、图片、session_id；保存上传图片", "标准化输入"],
            ["2", "意图解析", "抽取 intent、scene、style、colors、budget、categories、avoid", "结构化需求"],
            ["3", "记忆加载", "读取短期会话、长期偏好、语义记忆、图谱偏好边", "用户画像"],
            ["4", "元数据过滤", "先按品类、颜色、季节、场景做硬过滤或软加权", "过滤条件"],
            ["5", "多路召回", "文本向量、图片向量、标签召回、图谱邻居扩展", "候选商品池"],
            ["6", "融合排序", "RRF 合并多路结果，LLM/Reranker 复排", "Top-K 候选"],
            ["7", "图谱补全", "从候选商品扩展搭配、平替、同风格商品", "关系增强候选"],
            ["8", "搭配组合", "按类别组成 1-3 套方案并做约束检查", "穿搭方案"],
            ["9", "回答生成", "生成解释、商品卡片、取舍说明、下一步问题", "前端响应"],
            ["10", "记忆更新", "抽取新偏好，写业务表、向量库和图谱边", "更新用户画像"],
        ],
        [0.45, 1.15, 3.05, 1.9],
    )


def add_memory_design(doc: Document):
    add_heading(doc, "5. 用户记忆设计")
    add_paragraph(
        doc,
        "用户记忆分为短期记忆、长期偏好记忆、语义记忆和图谱记忆四层。一期必须实现前三层，并将稳定偏好同步为 Neo4j 中的 User-Style/User-Scene/User-Feature 关系边。",
    )
    add_table(
        doc,
        ["记忆层", "存储位置", "内容", "更新方式", "用途"],
        [
            ["短期记忆", "Redis + messages 表", "当前会话上下文、本轮约束、最近推荐结果", "每轮对话写入，过期淘汰", "多轮追问和上下文承接"],
            ["长期偏好记忆", "user_preferences 表", "颜色、风格、预算、禁忌、舒适度等稳定偏好", "LLM 抽取 + 用户显式确认", "个性化过滤和加权"],
            ["语义记忆", "user_preference_vectors", "用户偏好摘要文本向量", "偏好摘要变化后重写向量", "“按我平时喜欢的风格”语义召回"],
            ["图谱记忆", "Neo4j", "User-LIKES-Style、User-AVOIDS-Feature、User-PREFERS-Scene", "长期偏好变化后同步关系边", "图谱扩展、关系解释、个性化推荐"],
        ],
        [1.0, 1.45, 2.05, 1.65, 1.45],
    )
    add_table(
        doc,
        ["记忆更新规则", "说明"],
        [
            ["显式优先", "用户明确说“我喜欢/不喜欢/以后都推荐”时，权重高，直接进入长期偏好"],
            ["隐式谨慎", "点击、保存、追问只能作为弱信号，不能直接覆盖长期偏好"],
            ["冲突处理", "新偏好与旧偏好冲突时保留版本，降低旧偏好权重，不直接删除"],
            ["用户可控", "前端必须提供偏好查看、删除、清空入口，避免记忆污染"],
            ["可解释", "推荐时说明“根据你偏好黑白灰、低调通勤风”"],
        ],
        [1.45, 4.85],
    )


def add_retrieval_design(doc: Document):
    add_heading(doc, "6. 混合检索与准确性设计")
    add_paragraph(
        doc,
        "系统不能只依赖文字向量检索。服饰推荐涉及颜色、材质、场景、搭配关系和视觉相似度，必须采用“结构化过滤 + 多模态向量召回 + 图谱扩展 + 重排序”的组合策略。",
    )
    add_table(
        doc,
        ["召回方式", "输入", "适用场景", "输出"],
        [
            ["文本向量召回", "用户 query、意图描述、偏好摘要", "自然语言商品搜索", "语义相似商品"],
            ["图片向量召回", "上传图片或商品图片", "找相似款、同风格款", "视觉相似商品"],
            ["标签过滤", "颜色、风格、季节、场景、材质", "强约束和软约束过滤", "符合条件商品集合"],
            ["图谱扩展", "候选商品、用户偏好节点", "搭配、平替、相似、同场景推荐", "关系邻居商品"],
            ["RRF 融合", "多路召回结果", "减少不同分数体系不可比问题", "融合排序列表"],
            ["LLM/Reranker 复排", "候选商品 + 用户需求", "最终推荐前质量控制", "Top-K 推荐结果"],
        ],
        [1.15, 1.25, 2.0, 1.9],
    )
    add_table(
        doc,
        ["准确性保障", "实现方式"],
        [
            ["检索评估集", "人工标注 20-50 条查询，每条 3-5 个相关商品，计算 HitRate@5、MRR、Recall@10"],
            ["消融对比", "分别评估 text only、image only、hybrid、hybrid+graph、hybrid+graph+rerank"],
            ["约束检查", "最终回答前检查颜色、场景、类别、禁忌是否满足，不满足则重新组合或反问"],
            ["推荐日志", "记录候选、融合分、图谱扩展来源、最终选择原因，便于回放和调参"],
        ],
        [1.5, 4.8],
    )


def add_tables(doc: Document):
    add_heading(doc, "7. 业务数据表设计")
    tables = {
        "users 用户表": [
            ["id", "TEXT", "主键", "用户唯一标识"],
            ["nickname", "TEXT", "普通索引", "昵称，可为空"],
            ["created_at", "DATETIME", "普通索引", "创建时间"],
            ["updated_at", "DATETIME", "无", "更新时间"],
        ],
        "sessions 会话表": [
            ["id", "TEXT", "主键", "会话 ID"],
            ["user_id", "TEXT", "外键索引 users.id", "所属用户"],
            ["title", "TEXT", "无", "会话标题"],
            ["status", "TEXT", "普通索引", "active / archived"],
            ["created_at", "DATETIME", "普通索引", "创建时间"],
            ["updated_at", "DATETIME", "普通索引", "更新时间"],
        ],
        "messages 消息表": [
            ["id", "TEXT", "主键", "消息 ID"],
            ["session_id", "TEXT", "外键索引 sessions.id", "所属会话"],
            ["role", "TEXT", "普通索引", "user / assistant / system / tool"],
            ["content", "TEXT", "全文索引可选", "消息正文"],
            ["image_path", "TEXT", "无", "上传图片路径"],
            ["metadata_json", "TEXT", "无", "模型参数、检索过程、工具调用记录"],
            ["created_at", "DATETIME", "普通索引", "创建时间"],
        ],
        "products 商品表": [
            ["id", "INTEGER", "主键", "商品 ID"],
            ["source_text", "TEXT", "无", "原始数据集 text"],
            ["category", "TEXT", "普通索引", "outer / bottom / shoes 等"],
            ["brand", "TEXT", "普通索引", "品牌，可从 name 抽取"],
            ["name", "TEXT", "普通索引", "商品名称"],
            ["summary", "TEXT", "无", "原始英文摘要"],
            ["image_path", "TEXT", "唯一索引", "商品图片路径"],
            ["image_summary", "TEXT", "无", "视觉描述"],
            ["description", "TEXT", "无", "营销/推荐描述"],
            ["status", "TEXT", "普通索引", "active / inactive"],
            ["created_at", "DATETIME", "普通索引", "创建时间"],
            ["updated_at", "DATETIME", "无", "更新时间"],
        ],
        "product_tags 商品标签表": [
            ["id", "INTEGER", "主键", "标签记录 ID"],
            ["product_id", "INTEGER", "外键索引 products.id", "关联商品"],
            ["tag_type", "TEXT", "联合索引 tag_type, tag_value", "color / style / season / scene / material / feature"],
            ["tag_value", "TEXT", "联合索引 tag_type, tag_value", "标准标签值"],
            ["confidence", "REAL", "无", "置信度"],
            ["source", "TEXT", "普通索引", "vl_model / llm / manual"],
        ],
        "product_relations 商品关系表": [
            ["id", "TEXT", "主键", "关系 ID"],
            ["source_product_id", "INTEGER", "联合索引 source_product_id, relation_type", "起始商品"],
            ["target_product_id", "INTEGER", "普通索引", "目标商品"],
            ["relation_type", "TEXT", "联合索引 source_product_id, relation_type", "matches_with / similar_to / alternative_of"],
            ["score", "REAL", "普通索引", "关系强度"],
            ["reason", "TEXT", "无", "关系原因"],
            ["source", "TEXT", "普通索引", "rule / llm / user_behavior"],
            ["created_at", "DATETIME", "普通索引", "创建时间"],
        ],
        "user_preferences 用户偏好表": [
            ["id", "TEXT", "主键", "偏好 ID"],
            ["user_id", "TEXT", "联合索引 user_id, preference_type", "所属用户"],
            ["preference_type", "TEXT", "联合索引 user_id, preference_type", "preferred_color / avoid_style / budget 等"],
            ["preference_value", "TEXT", "联合索引 preference_type, preference_value", "偏好值"],
            ["weight", "REAL", "普通索引", "偏好强度"],
            ["confidence", "REAL", "无", "置信度"],
            ["source", "TEXT", "普通索引", "explicit / inferred / behavior"],
            ["source_message_id", "TEXT", "外键 messages.id", "来源消息"],
            ["created_at", "DATETIME", "普通索引", "创建时间"],
            ["updated_at", "DATETIME", "无", "更新时间"],
        ],
        "user_memory_summaries 用户记忆摘要表": [
            ["id", "TEXT", "主键", "摘要 ID"],
            ["user_id", "TEXT", "唯一索引", "所属用户"],
            ["summary", "TEXT", "无", "长期偏好自然语言摘要"],
            ["vector_doc_id", "TEXT", "普通索引", "对应 user_preference_vectors 文档 ID"],
            ["version", "INTEGER", "普通索引", "版本号"],
            ["updated_at", "DATETIME", "普通索引", "更新时间"],
        ],
        "recommendation_logs 推荐日志表": [
            ["id", "TEXT", "主键", "推荐日志 ID"],
            ["session_id", "TEXT", "外键索引 sessions.id", "所属会话"],
            ["user_id", "TEXT", "外键索引 users.id", "所属用户"],
            ["intent_json", "TEXT", "无", "结构化意图"],
            ["candidate_ids_json", "TEXT", "无", "候选商品"],
            ["graph_expansion_json", "TEXT", "无", "图谱扩展来源"],
            ["result_json", "TEXT", "无", "最终推荐方案"],
            ["strategy", "TEXT", "普通索引", "text / image / hybrid / graph / outfit"],
            ["latency_ms", "INTEGER", "普通索引", "耗时"],
            ["created_at", "DATETIME", "普通索引", "创建时间"],
        ],
        "feedback 用户反馈表": [
            ["id", "TEXT", "主键", "反馈 ID"],
            ["recommendation_id", "TEXT", "外键 recommendation_logs.id", "关联推荐"],
            ["user_id", "TEXT", "外键索引 users.id", "用户"],
            ["product_id", "INTEGER", "外键索引 products.id", "商品，可为空"],
            ["feedback_type", "TEXT", "普通索引", "like / dislike / click / save / reject"],
            ["comment", "TEXT", "无", "文字反馈"],
            ["created_at", "DATETIME", "普通索引", "创建时间"],
        ],
    }

    for title, rows in tables.items():
        add_heading(doc, title, level=2)
        add_table(doc, ["字段", "类型", "索引", "说明"], rows, [1.25, 1.0, 1.75, 2.5])


def add_graph_schema(doc: Document):
    add_heading(doc, "8. 图数据库设计（一期纳入）")
    add_paragraph(
        doc,
        "图数据库 Neo4j 在一期即纳入，核心价值是补足向量检索无法稳定表达的商品关系、搭配关系、平替关系、用户偏好关系和推荐解释路径。关系型表 product_relations 作为图谱的事实备份，Neo4j 作为在线关系查询和路径解释引擎。",
    )
    add_heading(doc, "8.1 节点设计", level=2)
    add_table(
        doc,
        ["节点标签", "主键", "核心属性", "用途"],
        [
            ["Product", "product_id", "name, category, brand, image_path, status", "商品节点"],
            ["Category", "name", "name, level", "品类层级"],
            ["Brand", "name", "name", "品牌关联"],
            ["Style", "name", "name, normalized_name", "风格，如 minimal / outdoor"],
            ["Scene", "name", "name, normalized_name", "场景，如 commute / dating / outdoor"],
            ["Color", "name", "name, hex 可选", "颜色约束和偏好"],
            ["Feature", "name", "name, type", "材质、版型、功能、禁忌特征"],
            ["User", "user_id", "user_id, created_at", "用户偏好图谱入口"],
        ],
        [1.15, 1.15, 2.4, 1.6],
    )
    add_heading(doc, "8.2 关系设计", level=2)
    add_table(
        doc,
        ["关系", "方向", "属性", "用途"],
        [
            ["BELONGS_TO", "Product -> Category", "source", "品类归属"],
            ["HAS_BRAND", "Product -> Brand", "source", "品牌关联"],
            ["HAS_STYLE", "Product -> Style", "score, source", "商品风格标签"],
            ["SUITABLE_FOR", "Product -> Scene", "score, reason, source", "适用场景"],
            ["HAS_COLOR", "Product -> Color", "score, source", "颜色标签"],
            ["HAS_FEATURE", "Product -> Feature", "score, source", "材质、功能、图案等特征"],
            ["MATCHES_WITH", "Product -> Product", "score, reason, source", "可搭配关系"],
            ["SIMILAR_TO", "Product -> Product", "score, reason, source", "相似款关系"],
            ["ALTERNATIVE_OF", "Product -> Product", "score, reason, source", "平替关系"],
            ["LIKES", "User -> Style/Color/Scene/Feature", "weight, confidence, source", "用户喜欢偏好"],
            ["AVOIDS", "User -> Style/Color/Feature", "weight, confidence, source", "用户禁忌偏好"],
            ["VIEWED/SAVED/REJECTED", "User -> Product", "count, last_at", "用户行为信号"],
        ],
        [1.25, 1.65, 1.75, 1.65],
    )
    add_heading(doc, "8.3 图谱查询示例", level=2)
    add_table(
        doc,
        ["业务问题", "图查询思路"],
        [
            ["这件外套适合搭什么裤子？", "Product-[:MATCHES_WITH]->Product，并过滤 target.category='bottom'"],
            ["有没有同风格平替？", "Product-[:ALTERNATIVE_OF|SIMILAR_TO]->Product，并结合价格/风格标签过滤"],
            ["按我的偏好推荐", "User-[:LIKES]->Style/Color/Scene 反向匹配 Product"],
            ["为什么推荐这件？", "返回 User 偏好节点到 Product 的关系路径作为解释依据"],
        ],
        [2.0, 4.3],
    )


def add_vector_schema(doc: Document):
    add_heading(doc, "9. 向量存储结构设计")
    add_table(
        doc,
        ["Collection", "模型", "维度", "Document", "Metadata", "用途"],
        [
            ["product_text_vectors", "text-embedding-v3", "1024", "description / image_summary / summary / normalized_tags", "product_id, field, category, colors, styles, scenes, seasons, graph_node_id", "文本语义召回"],
            ["product_image_vectors", "OpenCLIP / Chinese-CLIP", "512", "image_path", "product_id, category, name, colors, styles, graph_node_id", "视觉相似召回"],
            ["user_preference_vectors", "text-embedding-v3", "1024", "用户偏好摘要", "user_id, version, updated_at, graph_node_id", "长期偏好语义召回"],
            ["session_memory_vectors", "text-embedding-v3", "1024", "会话摘要、关键约束", "session_id, user_id, turn_range", "长会话记忆召回"],
        ],
        [1.35, 1.25, 0.5, 1.65, 2.05, 1.25],
    )


def add_langgraph_and_milestones(doc: Document):
    add_heading(doc, "10. LangGraph 编排设计")
    add_table(
        doc,
        ["节点", "职责", "循环/分支", "涉及存储"],
        [
            ["parse_intent", "解析意图、场景、预算、颜色、品类、禁忌", "失败降级为普通搜索", "无"],
            ["load_memory", "加载短期记忆、长期偏好、语义记忆和图谱偏好", "读取失败使用空画像", "Redis / DB / Vector / Neo4j"],
            ["decide_next_step", "判断反问、检索、对比、搭配或偏好更新", "条件分支", "无"],
            ["retrieve_candidates", "文本、图片、标签、图谱多路召回", "低质量时 query 改写后重试", "Chroma / Neo4j / DB"],
            ["expand_graph", "对候选商品做搭配、平替、相似关系扩展", "缺品类时补召回", "Neo4j"],
            ["rank_candidates", "RRF 融合 + Reranker 复排", "结果不足时回到召回", "Chroma / DB"],
            ["compose_outfit", "组合外套、上装、下装、鞋子方案", "不满足约束则重组", "DB / Neo4j"],
            ["check_constraints", "检查颜色、场景、预算、禁忌、品类完整性", "不通过则重试或反问", "无"],
            ["generate_answer", "生成推荐解释和商品卡片数据", "失败模板化降级", "DB"],
            ["update_memory", "更新偏好表、记忆向量和图谱边", "异步执行", "DB / Vector / Neo4j"],
        ],
        [1.3, 2.05, 1.5, 1.45],
    )

    add_heading(doc, "11. 里程碑规划")
    add_table(
        doc,
        ["版本", "目标", "交付内容", "验收标准"],
        [
            ["v0.3", "架构补强版", "FastAPI、Neo4j、用户记忆、混合检索、图谱扩展", "文本/图片/图文查询均可返回商品卡片和关系解释"],
            ["v0.4", "私人导购版", "偏好记忆面板、主动反问、用户画像、SSE 流式输出", "多轮对话能记住并应用用户偏好"],
            ["v0.5", "成套搭配版", "OutfitComposer、搭配图谱、约束检查、成套解释", "可稳定生成 1-3 套完整穿搭"],
            ["v0.6", "Web 演示版", "OpenClaw 风格低权限前端、商品详情侧栏、推荐日志回放", "可用于技术评审和产品演示"],
        ],
        [0.75, 1.3, 2.9, 1.35],
    )


def add_risks_and_conclusion(doc: Document):
    add_heading(doc, "12. 技术风险与应对")
    add_table(
        doc,
        ["风险", "影响", "应对策略"],
        [
            ["图谱关系质量不足", "搭配/平替推荐不可靠", "一期采用规则 + LLM 双来源，保留 score/source/reason，人工抽样校验"],
            ["用户记忆污染", "错误偏好影响推荐", "显式偏好高权重，隐式行为低权重，支持用户删除记忆"],
            ["只靠文本检索不准", "颜色、材质、视觉风格召回差", "结构化标签过滤 + 图片向量 + 图谱扩展 + Rerank"],
            ["模型输出不稳定", "JSON 解析失败、解释漂移", "JSON Schema 校验、重试、模板降级"],
            ["OpenCLIP 中文弱", "中文查图效果差", "中文 query 先翻译，二期评估 Chinese-CLIP"],
            ["前端越权", "用户通过对话触发危险操作", "低权限业务 API、工具白名单、禁止命令执行/文件编辑接口"],
        ],
        [1.55, 1.55, 3.15],
    )

    add_heading(doc, "13. 结论")
    add_paragraph(
        doc,
        "本计划书已将图数据库调整为一期核心组件。最终架构不是单纯文本 RAG，而是“多模态向量检索 + 结构化标签过滤 + Neo4j 商品关系图谱 + 用户记忆 + LangGraph 受控编排”的 AI 私人导购系统。",
    )
    add_paragraph(
        doc,
        "一期评审重点：先完成 Neo4j 图谱结构、用户记忆闭环、混合检索准确性验证和低权限前端接入，再扩展成套穿搭和商品对比能力。",
        "一期评审重点：",
    )


def add_footer(doc: Document):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("AI 私人导购应用项目计划书 · 技术评审版")
        set_run_font(run, size=9, color="666666")


def build():
    doc = setup_document()
    add_cover(doc)
    add_business_requirements(doc)
    add_layered_architecture(doc)
    add_tech_stack(doc)
    add_data_flow(doc)
    add_memory_design(doc)
    add_retrieval_design(doc)
    add_tables(doc)
    add_graph_schema(doc)
    add_vector_schema(doc)
    add_langgraph_and_milestones(doc)
    add_risks_and_conclusion(doc)
    add_footer(doc)
    try:
        doc.save(OUT_PATH)
        return OUT_PATH
    except PermissionError:
        # 原文件可能正在被 Word/WPS 打开，保存为 v2 版本避免中断。
        doc.save(FALLBACK_OUT_PATH)
        return FALLBACK_OUT_PATH


if __name__ == "__main__":
    print(build())
